# =========================================================
# SRMA Web App (Streamlit) - General-purpose SR/MA assistant
# Key design goals:
# - Minimal user input: default to P + I only for PubMed query generation
# - C / O are advanced optional inputs (do NOT affect search unless user enables strict mode)
# - Optional LLM (OpenAI-compatible) for feasibility + extraction
# - Optional python-docx / matplotlib / PyPDF2 (app will not crash if missing)
# - No credential storage: provide OpenURL/EZproxy links only
# =========================================================

from __future__ import annotations

import os
import io
import re
import math
import json
import time
import html
from typing import Dict, List, Tuple, Optional

import requests
import pandas as pd
import streamlit as st

# Optional: PDF text extraction
try:
    from PyPDF2 import PdfReader
    HAS_PYPDF2 = True
except Exception:
    HAS_PYPDF2 = False

# Optional: plotting
try:
    import matplotlib.pyplot as plt
    HAS_MPL = True
except Exception:
    HAS_MPL = False

# Optional: Word export (python-docx)
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

# --------------------- Page ---------------------
st.set_page_config(page_title="SRMA", layout="wide")

# --------------------- Basic UI CSS ---------------------
CSS = """
<style>
.card {
    border: 1px solid #dde2eb;
    border-radius: 10px;
    padding: 0.9rem 1rem;
    margin-bottom: 0.9rem;
    background-color: #ffffff;
}
.meta { font-size: 0.85rem; color: #444; }
.badge { display: inline-block; padding: 0.15rem 0.55rem; border-radius: 999px;
         font-size: 0.78rem; margin-right: 0.35rem; border: 1px solid rgba(0,0,0,0.06); }
.badge-include { background: #d1fae5; color: #065f46; }
.badge-exclude { background: #fee2e2; color: #991b1b; }
.badge-unsure  { background: #e0f2fe; color: #075985; }
.small { font-size: 0.85rem; color: #666; }
.kpi {
    border: 1px solid #e5e7eb;
    border-radius: 10px;
    padding: 0.75rem 0.9rem;
    background: #f9fafb;
}
.kpi .label { font-size: 0.8rem; color: #6b7280; }
.kpi .value { font-size: 1.2rem; font-weight: 700; color: #111827; }
hr.soft { border: none; border-top: 1px solid #eef2f7; margin: 0.8rem 0; }
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)

st.title("SRMA")

# =========================================================
# [A] Optional access gate (prevent public abuse)
# =========================================================
APP_PASSWORD = os.getenv("APP_PASSWORD", "").strip()
if APP_PASSWORD:
    with st.sidebar:
        st.subheader("Access")
        pw = st.text_input("App password", type="password")
    if pw != APP_PASSWORD:
        st.warning("Password required.")
        st.stop()

# =========================================================
# [0] Helpers
# =========================================================
def ensure_columns(df: pd.DataFrame, cols: List[str], default="") -> pd.DataFrame:
    for c in cols:
        if c not in df.columns:
            df[c] = default
    return df

def safe_int(x, default=0) -> int:
    try:
        return int(x)
    except Exception:
        return default

def norm_text(x: str) -> str:
    if not x:
        return ""
    x = html.unescape(str(x))
    x = re.sub(r"\s+", " ", x).strip()
    return x

def short(s: str, n=120) -> str:
    s = s or ""
    return (s[:n] + "…") if len(s) > n else s

def to_csv_bytes(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8-sig")

def doi_to_url(doi: str) -> str:
    doi = (doi or "").strip()
    return f"https://doi.org/{doi}" if doi else ""

def pubmed_url(pmid: str) -> str:
    pmid = (pmid or "").strip().replace("PMID:", "").strip()
    return f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/" if pmid else ""

def pmc_url(pmcid: str) -> str:
    pmcid = (pmcid or "").strip()
    if not pmcid:
        return ""
    if not pmcid.upper().startswith("PMC"):
        pmcid = "PMC" + pmcid
    return f"https://pmc.ncbi.nlm.nih.gov/articles/{pmcid}/"

def badge_html(label: str) -> str:
    label = label or "Unsure"
    if label == "Include":
        cls = "badge badge-include"
    elif label == "Exclude":
        cls = "badge badge-exclude"
    else:
        cls = "badge badge-unsure"
    return f'<span class="{cls}">{label}</span>'

def json_from_text(s: str) -> Optional[dict]:
    """Best-effort parse JSON object from a messy LLM response."""
    if not s:
        return None
    s = s.strip()
    try:
        return json.loads(s)
    except Exception:
        pass
    m = re.search(r"\{.*\}", s, flags=re.S)
    if m:
        try:
            return json.loads(m.group(0))
        except Exception:
            return None
    return None

# =========================================================
# [1] Institutional link helpers (resolver / ezproxy)
# =========================================================
def build_openurl(resolver_base: str, doi: str = "", pmid: str = "", title: str = "") -> str:
    resolver_base = (resolver_base or "").strip()
    if not resolver_base:
        return ""
    params = []
    params.append("url_ver=Z39.88-2004")
    params.append("ctx_ver=Z39.88-2004")
    if doi:
        params.append("rft_id=" + requests.utils.quote(f"info:doi/{doi}"))
    elif pmid:
        params.append("rft_id=" + requests.utils.quote(f"info:pmid/{pmid}"))
    if title:
        params.append("rft.title=" + requests.utils.quote(title[:200]))
    joiner = "&" if "?" in resolver_base else "?"
    return resolver_base + joiner + "&".join(params)

def apply_ezproxy(ezproxy_prefix: str, url: str) -> str:
    ezproxy_prefix = (ezproxy_prefix or "").strip()
    url = (url or "").strip()
    if not ezproxy_prefix or not url:
        return url
    if "url=" in ezproxy_prefix:
        return ezproxy_prefix + requests.utils.quote(url, safe="")
    if ezproxy_prefix.endswith("/"):
        ezproxy_prefix = ezproxy_prefix[:-1]
    return ezproxy_prefix + "/login?url=" + requests.utils.quote(url, safe="")

# =========================================================
# [2] MeSH suggestion sync (NLM MeSH lookup)
# =========================================================
@st.cache_data(show_spinner=False, ttl=60*60)
def mesh_suggest(term: str, limit: int = 6) -> List[str]:
    term = (term or "").strip()
    if not term:
        return []
    url = "https://id.nlm.nih.gov/mesh/lookup/descriptor"
    params = {"label": term, "match": "contains", "limit": str(limit)}
    try:
        r = requests.get(url, params=params, timeout=20)
        r.raise_for_status()
        data = r.json()
        labels = []
        for item in data:
            lab = item.get("label")
            if lab:
                labels.append(lab)
        seen, out = set(), []
        for x in labels:
            if x not in seen:
                out.append(x); seen.add(x)
        return out[:limit]
    except Exception:
        return []

def build_pubmed_block(term: str, mesh_label: Optional[str] = None) -> str:
    """
    Create a PubMed boolean block.
    If user already wrote advanced syntax (OR/AND/NOT/brackets/field tags), treat as raw block.
    """
    term = (term or "").strip()
    if not term:
        return ""
    low = term.lower()

    # Heuristic: if user typed boolean logic or field tags, do not rewrite.
    if any(x in low for x in [" or ", " and ", " not ", "[tiab]", "[mesh", "[mh]", "(", ")", "\""]):
        return term

    if mesh_label and mesh_label.strip():
        mesh_label = mesh_label.strip()
        return f'({term}[tiab] OR "{mesh_label}"[MeSH Terms])'
    return f'({term}[tiab] OR "{term}"[MeSH Terms])'

# =========================================================
# [3] PubMed / ClinicalTrials.gov fetchers
# =========================================================
NCBI_ESEARCH = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
NCBI_EFETCH  = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi"

def pubmed_esearch_ids(query: str, retstart: int, retmax: int) -> Tuple[List[str], int]:
    params = {"db": "pubmed", "term": query, "retmode": "json", "retstart": retstart, "retmax": retmax}
    r = requests.get(NCBI_ESEARCH, params=params, timeout=30)
    r.raise_for_status()
    js = r.json().get("esearchresult", {})
    ids = js.get("idlist", []) or []
    count = safe_int(js.get("count", 0), 0)
    return ids, count

def pubmed_efetch_xml(pmids: List[str]) -> str:
    params = {"db": "pubmed", "id": ",".join(pmids), "retmode": "xml"}
    r = requests.get(NCBI_EFETCH, params=params, timeout=90)
    r.raise_for_status()
    return r.text

def parse_pubmed_xml(xml_text: str) -> pd.DataFrame:
    import xml.etree.ElementTree as ET
    root = ET.fromstring(xml_text)
    rows = []
    for art in root.findall(".//PubmedArticle"):
        pmid = (art.findtext(".//PMID") or "").strip()
        title = norm_text(art.findtext(".//ArticleTitle") or "")
        ab_parts = []
        for ab in art.findall(".//AbstractText"):
            if ab.text:
                ab_parts.append(norm_text(ab.text))
        abstract = " ".join([x for x in ab_parts if x]).strip()
        year = (art.findtext(".//PubDate/Year") or "").strip()
        first_author = ""
        a0 = art.find(".//AuthorList/Author[1]")
        if a0 is not None:
            last = (a0.findtext("LastName") or "").strip()
            ini  = (a0.findtext("Initials") or "").strip()
            first_author = f"{last} {ini}".strip() if (last or ini) else ""
        journal = norm_text(art.findtext(".//Journal/Title") or "")
        doi = ""; pmcid = ""
        for aid in art.findall(".//ArticleIdList/ArticleId"):
            idt = (aid.get("IdType") or "").lower()
            val = (aid.text or "").strip()
            if idt == "doi" and val and not doi:
                doi = val
            if idt == "pmc" and val and not pmcid:
                pmcid = val
        rows.append({
            "record_id": f"PMID:{pmid}" if pmid else "",
            "pmid": pmid,
            "pmcid": pmcid,
            "doi": doi,
            "title": title,
            "abstract": abstract,
            "year": year,
            "first_author": first_author,
            "journal": journal,
            "source": "PubMed",
            "url": pubmed_url(pmid),
            "doi_url": doi_to_url(doi),
            "pmc_url": pmc_url(pmcid) if pmcid else "",
        })
    df = pd.DataFrame(rows)
    return ensure_columns(df, ["record_id","pmid","pmcid","doi","title","abstract","year","first_author","journal","source","url","doi_url","pmc_url"], default="")

def fetch_pubmed(query: str, max_records: int = 0, batch_size: int = 200, polite_delay: float = 0.0) -> Tuple[pd.DataFrame,int]:
    query = (query or "").strip()
    if not query:
        return pd.DataFrame(), 0

    ids, count = pubmed_esearch_ids(query, retstart=0, retmax=min(batch_size, 500))
    all_ids = list(ids)

    target = min(count, max_records) if (max_records and max_records > 0) else count
    while len(all_ids) < target:
        retstart = len(all_ids)
        need = min(batch_size, target - len(all_ids))
        ids, _ = pubmed_esearch_ids(query, retstart=retstart, retmax=need)
        if not ids:
            break
        all_ids.extend(ids)
        if polite_delay > 0:
            time.sleep(polite_delay)
        if "pubmed_progress" in st.session_state:
            st.session_state.pubmed_progress.progress(min(1.0, len(all_ids) / max(target, 1)))

    rows = []
    for i in range(0, len(all_ids), batch_size):
        chunk = all_ids[i:i+batch_size]
        xml = pubmed_efetch_xml(chunk)
        df = parse_pubmed_xml(xml)
        if not df.empty:
            rows.append(df)
        if polite_delay > 0:
            time.sleep(polite_delay)
        if "pubmed_fetch_progress" in st.session_state:
            st.session_state.pubmed_fetch_progress.progress(min(1.0, (i+len(chunk)) / max(len(all_ids), 1)))

    if not rows:
        return pd.DataFrame(), count
    out = pd.concat(rows, ignore_index=True)
    out = ensure_columns(out, ["record_id","pmid","pmcid","doi","title","abstract","year","first_author","journal","source","url","doi_url","pmc_url"], default="")
    return out, count

def fetch_ctgov(query: str, max_records: int = 200) -> pd.DataFrame:
    query = (query or "").strip()
    if not query or max_records <= 0:
        return pd.DataFrame()
    fields = ["NCTId","BriefTitle","BriefSummary","StartDate","StudyType","Condition"]
    url = "https://clinicaltrials.gov/api/query/study_fields"
    params = {"expr": query, "fields": ",".join(fields), "min_rnk": 1, "max_rnk": int(max_records), "fmt":"json"}
    r = requests.get(url, params=params, timeout=60)
    r.raise_for_status()
    studies = r.json().get("StudyFieldsResponse", {}).get("StudyFields", []) or []
    out = []
    for stf in studies:
        def first(k):
            v = stf.get(k, [])
            return v[0] if v else ""
        nct = first("NCTId")
        title = first("BriefTitle")
        abstract = first("BriefSummary")
        start = first("StartDate")
        year = start.split()[-1] if start else ""
        rid = f"NCT:{nct}" if nct else f"NCT:IDX:{len(out)}"
        url2 = f"https://clinicaltrials.gov/study/{nct}" if nct else ""
        out.append({
            "record_id": rid,
            "pmid": "",
            "pmcid": "",
            "doi": "",
            "title": norm_text(title),
            "abstract": norm_text(abstract),
            "year": year,
            "first_author": "",
            "journal": "",
            "source": "ClinicalTrials.gov",
            "url": url2,
            "doi_url": "",
            "pmc_url": "",
        })
    df = pd.DataFrame(out)
    return ensure_columns(df, ["record_id","pmid","pmcid","doi","title","abstract","year","first_author","journal","source","url","doi_url","pmc_url"], default="")

def deduplicate(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df["title_norm"] = df["title"].fillna("").str.lower().str.replace(r"\s+"," ", regex=True).str.strip()
    df["doi_norm"] = df["doi"].fillna("").str.lower().str.strip()
    if df["doi_norm"].astype(bool).any():
        df = df.sort_values(["doi_norm","source"]).drop_duplicates(subset=["doi_norm"], keep="first")
    df = df.sort_values(["title_norm","year","source"]).drop_duplicates(subset=["title_norm","year"], keep="first")
    df = df.drop(columns=["title_norm","doi_norm"], errors="ignore")
    return df.reset_index(drop=True)

# =========================================================
# [4] LLM (OpenAI-compatible) - optional
# =========================================================
def llm_available() -> bool:
    api_key = (st.session_state.get("LLM_API_KEY") or "").strip()
    base = (st.session_state.get("LLM_BASE_URL") or "").strip()
    model = (st.session_state.get("LLM_MODEL") or "").strip()
    return bool(api_key and base and model)

def llm_chat(messages: List[dict], temperature: float = 0.2, timeout: int = 90) -> Optional[str]:
    """
    OpenAI-compatible: POST {base}/v1/chat/completions
    """
    base = (st.session_state.get("LLM_BASE_URL") or "").strip().rstrip("/")
    api_key = (st.session_state.get("LLM_API_KEY") or "").strip()
    model = (st.session_state.get("LLM_MODEL") or "").strip()
    if not (base and api_key and model):
        return None
    url = base + "/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {"model": model, "messages": messages, "temperature": float(temperature)}
    try:
        r = requests.post(url, headers=headers, json=payload, timeout=timeout)
        r.raise_for_status()
        js = r.json()
        return js["choices"][0]["message"]["content"]
    except Exception as e:
        st.warning(f"LLM 呼叫失敗（將改用規則法/或跳過）：{e}")
        return None

# =========================================================
# [5] Rule-based AI screening (fallback, higher recall)
# =========================================================
def _count_hits(text_low: str, term: str) -> int:
    term = (term or "").strip()
    if not term:
        return 0
    toks = [t.strip().lower() for t in re.split(r"[,\s]+", term) if t.strip()]
    return sum(1 for t in toks if t and t in text_low)

def ai_screen_rule_based(row: pd.Series, pico: Dict[str,str]) -> Dict:
    """
    High-recall rule-based TA screening:
    - Prefer to keep candidates (Include/Unsure) for later FT over-aggressive Exclude.
    - Default uses P + I only (C/O not required).
    """
    title = row.get("title","") or ""
    abstract = row.get("abstract","") or ""
    text = (title + " " + abstract).lower()

    P = pico.get("P","") or ""
    I = pico.get("I","") or ""
    X = pico.get("X","") or ""

    p_hit = _count_hits(text, P)
    i_hit = _count_hits(text, I)
    x_hit = _count_hits(text, X) if X else 0

    is_trial = any(w in text for w in ["randomized", "randomised", "trial", "controlled", "prospective", "double-blind", "single-blind"])
    is_basic = any(w in text for w in ["in vitro", "cell line", "mouse", "mice", "rat", "animal model", "transcriptome", "proteomic", "pbmc"])
    is_case_report = any(w in text for w in ["case report", "case series"])

    # Hard excludes
    if X.strip() and x_hit:
        label = "Exclude"
        reason = "NOT keyword hit"
        conf = 0.9
        return {"label": label, "reason": reason, "confidence": conf, "rule_trace": "NOT"}

    if is_basic and not is_trial:
        label = "Exclude"
        reason = "Basic research likely (and not trial-like)"
        conf = 0.8
        return {"label": label, "reason": reason, "confidence": conf, "rule_trace": "basic"}

    if is_case_report and not is_trial:
        label = "Exclude"
        reason = "Case report/series likely (and not trial-like)"
        conf = 0.8
        return {"label": label, "reason": reason, "confidence": conf, "rule_trace": "case"}

    # Include if P+I both appear OR trial-like + one appears
    if (P.strip() and I.strip() and p_hit and i_hit) and (is_trial or len(text) > 0):
        label = "Include"
        reason = "P+I hit; candidate for full-text"
        conf = 0.75
        return {"label": label, "reason": reason, "confidence": conf, "rule_trace": "PI"}

    if is_trial and ((P.strip() and p_hit) or (I.strip() and i_hit)):
        label = "Include"
        reason = "Trial-like + P or I hit; keep for FT"
        conf = 0.7
        return {"label": label, "reason": reason, "confidence": conf, "rule_trace": "trial+P/I"}

    # Otherwise: Unsure (do not over-exclude)
    label = "Unsure"
    reason = "Insufficient TA evidence; keep as unsure"
    conf = 0.4
    return {"label": label, "reason": reason, "confidence": conf, "rule_trace": "unsure"}

# =========================================================
# [6] Feasibility scan + optional LLM feasibility report
# =========================================================
def quick_srma_scan(pubmed_query: str, top_n: int = 20) -> Dict:
    q = (pubmed_query or "").strip()
    if not q:
        return {"summary":"(未提供 query)", "hits": pd.DataFrame()}
    sr_filter = '(systematic review[Publication Type] OR meta-analysis[Publication Type] OR "systematic review"[tiab] OR "meta analysis"[tiab] OR "network meta-analysis"[tiab])'
    q_sr = f"({q}) AND {sr_filter}"
    try:
        ids, count = pubmed_esearch_ids(q_sr, 0, min(top_n, 50))
        df_hits = pd.DataFrame()
        if ids:
            xml = pubmed_efetch_xml(ids)
            df_hits = parse_pubmed_xml(xml)[["pmid","doi","pmcid","title","year","first_author","journal","url","doi_url","pmc_url"]]
        summary = f"PubMed 疑似 SR/MA：count ≈ {count}；列出前 {min(len(ids), top_n)} 篇。"
        return {"summary": summary, "hits": df_hits}
    except Exception as e:
        return {"summary": f"SR/MA 掃描失敗：{e}", "hits": pd.DataFrame()}

def feasibility_report_llm(pico: Dict[str,str], goal_mode: str, clinical_context: str,
                           srma_hits: pd.DataFrame, current_query: str) -> Dict:
    sr_list = []
    if isinstance(srma_hits, pd.DataFrame) and not srma_hits.empty:
        for _, r in srma_hits.head(15).iterrows():
            sr_list.append({"pmid": r.get("pmid",""), "year": r.get("year",""), "title": r.get("title","")})
    sys = "You are a senior SR/MA methodologist. Produce rigorous, auditable outputs."
    user = {
        "task": "Stage A feasibility + topic scoping + criteria + extraction schema planning",
        "goal_mode": goal_mode,
        "clinical_context": clinical_context,
        "input_pico": pico,
        "current_pubmed_query": current_query,
        "known_srma_hits": sr_list,
        "requirements": [
            "Draft inclusion/exclusion criteria as executable rules with IDs.",
            "Propose dynamic extraction schema based on topic and typical RCT outcomes.",
            "Suggest search strategy with MeSH + synonyms; ensure it is feasible and publishable.",
            "Output strictly as JSON with keys: refined_pico, inclusion_rules, exclusion_rules, search_strategy, extraction_schema, rationale."
        ]
    }
    messages = [{"role": "system", "content": sys},
                {"role": "user", "content": json.dumps(user, ensure_ascii=False)}]
    txt = llm_chat(messages, temperature=0.2, timeout=120)
    js = json_from_text(txt or "")
    if not js:
        return {"error": "LLM 回傳無法解析為 JSON。", "raw": txt or ""}
    return js

# =========================================================
# [7] Extraction prompt template (OCR/table/figure aware)
# =========================================================
def build_extraction_prompt(schema: dict, pico: dict, criteria_text: str) -> str:
    base_cols = schema.get("base_cols", [])
    outcomes = schema.get("outcomes", [])
    prompt = f"""
[Task E] Full-text review + data extraction (SR/MA)

OCR/figure/table instructions:
1) If PDF text is empty or looks scanned (no selectable text), state "OCR REQUIRED" and instruct OCR, then proceed best-effort.
2) Prioritize tables, figure legends, appendix/supplement.
3) When extracting numeric values, include source location: page if known; otherwise table/figure number or section heading.

Inclusion/Exclusion criteria (for traceability):
{criteria_text if criteria_text else "(No formal criteria provided; use P+I as primary guidance.)"}

Inputs:
P={pico.get("P","")}
I={pico.get("I","")}
C={pico.get("C","")}
O={pico.get("O","")}
NOT={pico.get("X","")}

Extraction schema (do not invent; leave blank if not found):
Base fields:
{json.dumps(base_cols, ensure_ascii=False)}

Outcome fields:
{json.dumps(outcomes, ensure_ascii=False)}

Also extract effect estimate if possible:
- effect_measure: OR/RR/HR/MD/SMD/RD
- effect, lower_CI, upper_CI, timepoint, unit
If not possible, leave blank and state why.

Output format: JSON with keys:
- fulltext_decision: "Include for meta-analysis" / "Exclude after full-text" / "Not reviewed yet"
- fulltext_reason
- extracted_fields: object mapping each schema field -> value (string)
- meta: effect_measure/effect/lower_CI/upper_CI/timepoint/unit
- missing_notes: list of what cannot be extracted and where to check (table/figure/page)
"""
    return prompt.strip()

# =========================================================
# [8] Session state
# =========================================================
def init_state():
    ss = st.session_state
    ss.setdefault("records_df", pd.DataFrame())
    ss.setdefault("pubmed_total_count", 0)
    ss.setdefault("ai_ta_df", pd.DataFrame())
    ss.setdefault("ta_final", {})
    ss.setdefault("criteria_text", "")
    ss.setdefault("feasibility", {"summary":"", "hits": pd.DataFrame(), "llm_json": None})
    ss.setdefault("ft_decisions", {})
    ss.setdefault("ft_reasons", {})
    ss.setdefault("ft_text", {})
    ss.setdefault("ft_note", {})
    ss.setdefault("rob2", {})
    ss.setdefault("extract_wide_df", pd.DataFrame())
    ss.setdefault("screen_cursor", 0)
    ss.setdefault("AUTO_WORD_BYTES", b"")
    ss.setdefault("AUTO_PRISMA", {})

    # Search building controls
    ss.setdefault("AUTO_FOLLOW", True)     # auto follow PICO/MeSH for query
    ss.setdefault("STRICT_CO", False)      # whether to include C/O in query
init_state()

# =========================================================
# [9] Sidebar settings (LLM + institutional resolver)
# =========================================================
with st.sidebar:
    st.subheader("設定")
    st.markdown("**全文導向（不存帳密）**")
    resolver_base = st.text_input(
        "OpenURL / Link resolver base（例：https://resolver.xxx.edu/openurl?）",
        value=st.session_state.get("RESOLVER_BASE",""),
        help="填機構 link resolver。每篇文獻會產生『全文連結』，使用者點了再自行登入下載。"
    )
    ezproxy_prefix = st.text_input(
        "EZproxy prefix（可選；例：https://ezproxy.xxx.edu/login?url=）",
        value=st.session_state.get("EZPROXY",""),
        help="若有 EZproxy，可填前綴。會把外部連結轉成 EZproxy 版本。"
    )
    st.session_state["RESOLVER_BASE"] = resolver_base
    st.session_state["EZPROXY"] = ezproxy_prefix

    st.markdown("---")
    st.markdown("**LLM（可選；OpenAI-compatible）**")
    llm_base = st.text_input("Base URL（例：https://api.openai.com）", value=st.session_state.get("LLM_BASE_URL",""))
    llm_key  = st.text_input("API Key", value=st.session_state.get("LLM_API_KEY",""), type="password")
    llm_model= st.text_input("Model", value=st.session_state.get("LLM_MODEL",""))
    st.session_state["LLM_BASE_URL"] = llm_base
    st.session_state["LLM_API_KEY"] = llm_key
    st.session_state["LLM_MODEL"] = llm_model

    if llm_available():
        st.success("LLM 已設定：可用於可行性報告/criteria/schema/更準確篩選/抽取")
    else:
        st.info("未設定 LLM：將使用規則法（穩定、可跑大量）；5B 抽取將不可用")

    st.markdown("---")
    if not HAS_DOCX:
        st.warning("目前環境未安裝 python-docx：Word 匯出功能停用（App 不會因此掛掉）。")
    if not HAS_MPL:
        st.warning("目前環境未安裝 matplotlib：森林圖功能停用。")
    if not HAS_PYPDF2:
        st.warning("目前環境未安裝 PyPDF2：PDF 抽字功能停用（可改用貼全文）。")

# =========================================================
# [10] Step 1: PICO + MeSH + Query (P+I default)
# =========================================================
st.header("Step 1. 輸入 PICO（預設只需 P+I）+ MeSH 同步 + NOT")

goal_mode = st.selectbox(
    "Goal mode",
    options=["Auto", "Fast / feasible", "Rigorous / comprehensive"],
    index=["Auto","Fast / feasible","Rigorous / comprehensive"].index(st.session_state.get("GOAL","Auto"))
)
st.session_state["GOAL"] = goal_mode

clinical_context = st.text_area(
    "Clinical context / scenario（可留空；不會自動進搜尋式）",
    value=st.session_state.get("CTX",""),
    height=80
)
st.session_state["CTX"] = clinical_context

# ---- Basic: P + I ----
colA, colB = st.columns(2)
with colA:
    P = st.text_input("P (Population / Topic) 〔建議必填〕", value=st.session_state.get("P",""))
with colB:
    I = st.text_input("I (Intervention / Exposure) 〔建議必填〕", value=st.session_state.get("I",""))

# ---- Advanced: C / O optional ----
with st.expander("進階（可選）：C / O（預設不影響 PubMed 搜尋式）", expanded=False):
    colC1, colC2 = st.columns(2)
    with colC1:
        C = st.text_input("C (Comparison) — 可留白", value=st.session_state.get("C",""))
    with colC2:
        O = st.text_input("O (Outcome) — 可留白", value=st.session_state.get("O",""))

strict_include_CO = st.checkbox(
    "嚴格把 C / O 納入 PubMed 檢索（會降低召回率；僅在你確定 TA 會寫出 C/O 時才建議開）",
    value=st.session_state.get("STRICT_CO", False)
)
st.session_state["STRICT_CO"] = strict_include_CO

exclude_not = st.text_input("排除關鍵字（NOT；例：animal OR pediatric OR case report）", value=st.session_state.get("X",""))
extra_kw = st.text_input("額外關鍵字（可留空）", value=st.session_state.get("EXTRA",""))
add_rct = st.checkbox("自動加上 RCT 關鍵字", value=st.session_state.get("ADD_RCT", True))

st.session_state["P"]=P
st.session_state["I"]=I
st.session_state["C"]=C
st.session_state["O"]=O
st.session_state["X"]=exclude_not
st.session_state["EXTRA"]=extra_kw
st.session_state["ADD_RCT"]=add_rct

with st.expander("MeSH term 同步（建議，可留空不用）", expanded=False):
    def mesh_picker(label: str, term: str, key_prefix: str) -> str:
        term = (term or "").strip()
        if not term:
            st.caption(f"{label}: (空白)")
            st.session_state[f"{key_prefix}_MESH"] = ""
            return ""
        sug = mesh_suggest(term)
        default = st.session_state.get(f"{key_prefix}_MESH","")
        choice = st.selectbox(
            f"{label} 的 MeSH 建議（可留空不用）",
            options=[""] + sug,
            index=([""]+sug).index(default) if default in ([""]+sug) else 0,
            key=f"{key_prefix}_mesh_select"
        )
        st.session_state[f"{key_prefix}_MESH"] = choice
        if sug:
            st.caption("建議：" + " / ".join(sug[:6]))
        else:
            st.caption("查不到建議（或 API 暫時不可用）。")
        return choice

    mesh_picker("P", P, "P")
    mesh_picker("I", I, "I")
    mesh_picker("C", C, "C")
    mesh_picker("O", O, "O")
    mesh_picker("Extra", extra_kw, "EXTRA")

# ---- Build PubMed query (DEFAULT: P + I only) ----
parts = []
if P.strip():
    parts.append(build_pubmed_block(P, st.session_state.get("P_MESH","")))
if I.strip():
    parts.append(build_pubmed_block(I, st.session_state.get("I_MESH","")))

# Only include C/O if strict mode enabled
if strict_include_CO:
    if C.strip():
        parts.append(build_pubmed_block(C, st.session_state.get("C_MESH","")))
    if O.strip():
        parts.append(build_pubmed_block(O, st.session_state.get("O_MESH","")))

if extra_kw.strip():
    parts.append(build_pubmed_block(extra_kw, st.session_state.get("EXTRA_MESH","")))

if add_rct:
    parts.append('(randomized controlled trial[tiab] OR randomised[tiab] OR randomized[tiab] OR trial[tiab] OR "Randomized Controlled Trial"[Publication Type])')

base_query = " AND ".join([p for p in parts if p]).strip()
pubmed_query_auto = ""
if base_query:
    pubmed_query_auto = f"({base_query})"
    if exclude_not.strip():
        pubmed_query_auto += f" NOT ({exclude_not.strip()})"

auto_follow = st.checkbox(
    "PubMed 搜尋式自動跟隨 PICO/MeSH（會覆蓋手動修改）",
    value=st.session_state.get("AUTO_FOLLOW", True)
)
st.session_state["AUTO_FOLLOW"] = auto_follow

has_min_inputs = bool(P.strip() or I.strip() or extra_kw.strip())

# Auto-follow: only override when there is meaningful auto query
if auto_follow:
    if has_min_inputs and pubmed_query_auto:
        st.session_state["PUBMED_QUERY"] = pubmed_query_auto
    elif not has_min_inputs:
        st.info("提示：請至少填 P 或 I（或 Extra），系統才會自動生成 PubMed 搜尋式。")
    else:
        # has_min_inputs but auto query empty -> do not wipe user's query
        st.warning("自動組 query 失敗（可能輸入為空或被視為無效）；不會覆蓋你原本的搜尋式。")

colQ1, colQ2 = st.columns([1, 3])
with colQ1:
    if st.button("套用：重建自動搜尋式"):
        if has_min_inputs and pubmed_query_auto:
            st.session_state["PUBMED_QUERY"] = pubmed_query_auto
        else:
            st.warning("無法重建：請先填 P 或 I（或 Extra）。")

with colQ2:
    pubmed_query = st.text_area("PubMed 搜尋式（可直接手動改）", value=st.session_state.get("PUBMED_QUERY",""), height=120)
    st.session_state["PUBMED_QUERY"] = pubmed_query

# =========================================================
# Step 1-1: extraction schema
# =========================================================
st.subheader("Step 1-1. extraction schema（欄位不寫死，可自訂；也可由 AI 建議）")

default_outcomes = st.session_state.get("OUTCOME_LINES", "Primary outcome\nSecondary outcome 1\nSecondary outcome 2")
outcome_lines = st.text_area("Outcome / 欄位名稱（每行一個，可自訂）", value=default_outcomes, height=120)
st.session_state["OUTCOME_LINES"] = outcome_lines

default_base_cols = st.session_state.get(
    "BASECOLS",
    "\n".join([
        "First author","Year","Country",
        "Intervention","Sample size (Intervention)",
        "Comparator","Sample size (Comparator)",
        "Follow-up","Key outcomes",
        "PICO (with NOT/exclude)","AI TA decision + reason",
        "Fulltext availability / note",
    ])
)
base_cols_text = st.text_area("基本欄位（每行一個）", value=default_base_cols, height=160)
st.session_state["BASECOLS"] = base_cols_text

schema = {
    "base_cols": [x.strip() for x in (base_cols_text or "").splitlines() if x.strip()],
    "outcomes": [x.strip() for x in (outcome_lines or "").splitlines() if x.strip()],
}

# =========================================================
# Step 1-2: Feasibility scan + optional LLM feasibility report
# =========================================================
st.header("Step 1-2. 可行性掃描（先找 SR/MA + criteria/schema 建議）")

colF1, colF2 = st.columns([1,1])
with colF1:
    if st.button("執行 SR/MA 掃描（PubMed）"):
        with st.spinner("掃描中…"):
            fb = quick_srma_scan(st.session_state.get("PUBMED_QUERY",""), top_n=20)
            st.session_state["feasibility"] = {"summary": fb["summary"], "hits": fb["hits"], "llm_json": None}

with colF2:
    st.caption("目的：確認是否已有 SR/MA，並為 criteria / schema 提供依據。")

fb = st.session_state.get("feasibility", {"summary":"", "hits": pd.DataFrame(), "llm_json": None})
if fb.get("summary"):
    st.info(fb["summary"])
    if isinstance(fb.get("hits"), pd.DataFrame) and not fb["hits"].empty:
        st.dataframe(fb["hits"], use_container_width=True)
        st.download_button("下載 SR/MA 清單（CSV）", data=to_csv_bytes(fb["hits"]),
                           file_name="feasibility_srma_hits.csv", mime="text/csv")

st.markdown("#### Inclusion / Exclusion criteria（可先手寫；也可由 LLM 生成）")
criteria_text = st.text_area("criteria（可留空）", value=st.session_state.get("criteria_text",""), height=220)
st.session_state["criteria_text"] = criteria_text

if llm_available():
    if st.button("用 LLM 產出：criteria + schema 建議"):
        with st.spinner("LLM 分析中…"):
            pico0 = {"P":P, "I":I, "C":C, "O":O, "X":exclude_not}
            js = feasibility_report_llm(pico0, goal_mode, clinical_context, fb.get("hits", pd.DataFrame()), st.session_state.get("PUBMED_QUERY",""))
            st.session_state["feasibility"]["llm_json"] = js

    js = st.session_state["feasibility"].get("llm_json")
    if js:
        if js.get("error"):
            st.error(js.get("error"))
            st.code(js.get("raw",""), language="text")
        else:
            st.subheader("LLM 建議（可套用到 criteria / schema）")
            st.code(json.dumps(js, ensure_ascii=False, indent=2), language="json")

            a1,a2 = st.columns(2)
            with a1:
                if st.button("套用：criteria"):
                    inc = js.get("inclusion_rules") or []
                    exc = js.get("exclusion_rules") or []
                    lines = ["Inclusion criteria (rule IDs):"]
                    for r in inc:
                        lines.append(f"- [{r.get('id','')}] {r.get('rule','')}")
                    lines.append("")
                    lines.append("Exclusion criteria (rule IDs):")
                    for r in exc:
                        lines.append(f"- [{r.get('id','')}] {r.get('rule','')}")
                    st.session_state["criteria_text"] = "\n".join(lines).strip()
                    st.rerun()
            with a2:
                if st.button("套用：extraction schema"):
                    es = js.get("extraction_schema") or {}
                    base_cols2 = es.get("base_cols") or schema["base_cols"]
                    outcomes2 = es.get("outcomes") or schema["outcomes"]
                    st.session_state["BASECOLS"] = "\n".join(base_cols2)
                    st.session_state["OUTCOME_LINES"] = "\n".join(outcomes2)
                    st.rerun()
else:
    st.caption("未設定 LLM：這裡提供 SR/MA 掃描 + 你手動輸入 criteria/schema。")

st.markdown("---")

# =========================================================
# [AUTO] Step 0: One-click pipeline + Word export
# =========================================================
st.header("Step 0. 一鍵自動跑（抓文獻→AI 粗篩→PRISMA→匯出）")

auto_use_ctgov = st.checkbox("一鍵跑時同時抓 ClinicalTrials.gov", value=False)
auto_pubmed_max = st.number_input("一鍵跑：PubMed 抓取上限（0=全部；太大會慢）", min_value=0, max_value=2000000, value=0, step=200)
auto_ctgov_max = st.number_input("一鍵跑：CT.gov 抓取上限", min_value=0, max_value=20000, value=200, step=50)
auto_batch = st.number_input("一鍵跑：AI TA batch size", min_value=20, max_value=500, value=200, step=20)
auto_engine = st.selectbox("一鍵跑：篩選引擎", options=["Auto (LLM if available else rule)", "Rule-based", "LLM"], index=0)

def ta_prompt(criteria_text_: str, pico_: dict, rec: dict) -> List[dict]:
    sys = "You are an SR/MA screening assistant. Be conservative. Output ONLY valid JSON."
    user = {
        "task": "Title/Abstract screening",
        "criteria_text": criteria_text_,
        "pico": pico_,
        "record": {"title": rec.get("title",""), "abstract": rec.get("abstract",""),
                   "year": rec.get("year",""), "source": rec.get("source","")},
        "output_schema": {"label": "Include/Exclude/Unsure", "confidence": "0..1",
                          "reason": "short rationale", "matched_rules": "list of rule IDs if any"},
        "policy": [
            "If insufficient info, choose Unsure.",
            "Prefer recall over precision in early screening."
        ]
    }
    return [{"role":"system","content":sys},{"role":"user","content":json.dumps(user, ensure_ascii=False)}]

def run_ta_llm(df_chunk: pd.DataFrame, pico_: dict, criteria_text_: str) -> pd.DataFrame:
    out = []
    for _, r in df_chunk.iterrows():
        rid = r["record_id"]
        rec = r.to_dict()
        txt = llm_chat(ta_prompt(criteria_text_, pico_, rec), temperature=0.1, timeout=90)
        js = json_from_text(txt or "")
        if not js:
            rb = ai_screen_rule_based(r, pico_)
            js = {"label": rb["label"], "confidence": rb["confidence"], "reason": rb["reason"], "matched_rules": []}
        out.append({
            "record_id": rid,
            "AI_label": js.get("label","Unsure"),
            "AI_confidence": js.get("confidence", 0.0),
            "AI_reason": js.get("reason",""),
            "AI_matched_rules": ",".join(js.get("matched_rules", []) or []),
        })
    return pd.DataFrame(out)

def run_ta_rule(df_chunk: pd.DataFrame, pico_: dict) -> pd.DataFrame:
    out = []
    for _, r in df_chunk.iterrows():
        rid = r["record_id"]
        rb = ai_screen_rule_based(r, pico_)
        out.append({
            "record_id": rid,
            "AI_label": rb["label"],
            "AI_confidence": rb["confidence"],
            "AI_reason": rb["reason"],
            "AI_matched_rules": rb.get("rule_trace",""),
        })
    return pd.DataFrame(out)

def build_summary_df(df_records_: pd.DataFrame) -> pd.DataFrame:
    ai = st.session_state.get("ai_ta_df", pd.DataFrame())
    if ai is None or ai.empty:
        ai = pd.DataFrame(columns=["record_id","AI_label","AI_confidence","AI_reason","AI_matched_rules"])
    rows = []
    for _, r in df_records_.iterrows():
        rid = r["record_id"]
        a = ai[ai["record_id"]==rid].head(1)
        a = a.iloc[0].to_dict() if not a.empty else {}
        openurl = build_openurl(st.session_state.get("RESOLVER_BASE",""),
                                doi=r.get("doi",""), pmid=r.get("pmid",""), title=r.get("title",""))
        openurl = apply_ezproxy(st.session_state.get("EZPROXY",""), openurl) if openurl else ""
        rows.append({
            "record_id": rid,
            "pmid": r.get("pmid",""),
            "doi": r.get("doi",""),
            "year": r.get("year",""),
            "first_author": r.get("first_author",""),
            "title": r.get("title",""),
            "source": r.get("source",""),
            "institution_openurl": openurl,
            "AI_label": a.get("AI_label",""),
            "AI_confidence": a.get("AI_confidence",""),
            "AI_reason": a.get("AI_reason",""),
            "AI_matched_rules": a.get("AI_matched_rules",""),
            "TA_final_AI": st.session_state["ta_final"].get(rid, "Unsure"),
            "FT_decision": st.session_state["ft_decisions"].get(rid, "Not reviewed yet"),
            "FT_reason": st.session_state["ft_reasons"].get(rid, ""),
            "FT_note_cannot_obtain": st.session_state["ft_note"].get(rid, ""),
            "FT_fulltext_text": st.session_state["ft_text"].get(rid, ""),
        })
    return pd.DataFrame(rows)

def compute_prisma_counts(df_all_raw: pd.DataFrame, df_dedup: pd.DataFrame) -> dict:
    identified = len(df_all_raw)
    deduped = len(df_dedup)
    dup_removed = max(0, identified - deduped)

    ta_vals = [st.session_state["ta_final"].get(rid, "Unsure") for rid in df_dedup["record_id"].tolist()]
    k_include = sum(1 for x in ta_vals if x == "Include")
    k_exclude = sum(1 for x in ta_vals if x == "Exclude")
    k_unsure  = sum(1 for x in ta_vals if x == "Unsure")

    ft_assessed = k_include + k_unsure
    ft = st.session_state.get("ft_decisions", {})
    inc_ft = sum(1 for _,v in ft.items() if v == "Include for meta-analysis")
    exc_ft = sum(1 for _,v in ft.items() if v == "Exclude after full-text")

    return {
        "identified": identified,
        "duplicates_removed": dup_removed,
        "records_screened": deduped,
        "records_excluded_ta": k_exclude,
        "fulltext_assessed": ft_assessed,
        "fulltext_excluded": exc_ft,
        "studies_included_meta": inc_ft,
        "ta_include": k_include,
        "ta_unsure": k_unsure,
    }

def auto_build_search_rows(pubmed_query_: str, pubmed_count: int, ctgov_used: bool, ctgov_count: int) -> List[List[str]]:
    rows = []
    rows.append(["PubMed", "1", pubmed_query_, str(pubmed_count)])
    if ctgov_used:
        rows.append(["ClinicalTrials.gov", "1", pubmed_query_, str(ctgov_count)])
    rows.append(["Embase", "", "(optional) translate from PubMed query", ""])
    rows.append(["OVID-Medline", "", "(optional) translate from PubMed query", ""])
    rows.append(["CENTRAL", "", "(optional) translate from PubMed query", ""])
    return rows

# ----- docx helpers (only used when HAS_DOCX=True) -----
def docx_add_heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12 if level==1 else 11)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def docx_add_table(doc, headers: List[str], rows: List[List[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i,h in enumerate(headers):
        hdr_cells[i].text = h
    for r in rows:
        cells = table.add_row().cells
        for i,val in enumerate(r):
            cells[i].text = str(val)
    return table

def make_word_supplement(pubmed_query_: str, prisma: dict, summary_df_: pd.DataFrame,
                         search_rows: List[List[str]]) -> bytes:
    if not HAS_DOCX:
        return b""

    doc = Document()
    title = doc.add_paragraph()
    tr = title.add_run("Supplement (auto-generated)")
    tr.bold = True
    tr.font.size = Pt(14)
    doc.add_paragraph(f"Generated time: {time.strftime('%Y-%m-%d %H:%M:%S')}")

    docx_add_heading(doc, "Supplement Table 1. Search Strategy", level=1)
    doc.add_paragraph("PubMed query:")
    doc.add_paragraph(pubmed_query_)

    headers = ["Database", "#", "Search syntax", "Citations found"]
    docx_add_table(doc, headers, search_rows)

    doc.add_paragraph()
    docx_add_heading(doc, "Supplement Figure 1. PRISMA Flow (numbers)", level=1)
    lines = [
        f"Records identified: {prisma.get('identified',0)}",
        f"Duplicates removed: {prisma.get('duplicates_removed',0)}",
        f"Records screened (deduplicated): {prisma.get('records_screened',0)}",
        f"Records excluded by title/abstract (AI): {prisma.get('records_excluded_ta',0)}",
        f"Full-text articles queued/assessed: {prisma.get('fulltext_assessed',0)}",
        f"Full-text excluded: {prisma.get('fulltext_excluded',0)}",
        f"Studies included in meta-analysis: {prisma.get('studies_included_meta',0)}",
    ]
    for s in lines:
        doc.add_paragraph(s)

    doc.add_paragraph()
    docx_add_heading(doc, "Screening snapshot (top 50)", level=1)
    snap = summary_df_.copy()
    keep = ["record_id","year","first_author","title","source","TA_final_AI","AI_reason","institution_openurl"]
    for c in keep:
        if c not in snap.columns:
            snap[c] = ""
    snap = snap[keep].head(50)
    rows = [[str(x) for x in r] for r in snap.values.tolist()]
    docx_add_table(doc, keep, rows)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def run_one_click_pipeline(auto_use_ctgov: bool, auto_pubmed_max: int, auto_ctgov_max: int,
                           auto_batch: int, auto_engine: str):
    q = (st.session_state.get("PUBMED_QUERY","") or "").strip()
    if not q:
        st.error("PubMed 搜尋式為空：請至少填 P 或 I（或手動貼搜尋式）。")
        return

    pico0 = {
        "P": st.session_state.get("P",""),
        "I": st.session_state.get("I",""),
        "C": st.session_state.get("C",""),
        "O": st.session_state.get("O",""),
        "X": st.session_state.get("X",""),
    }
    criteria_now = st.session_state.get("criteria_text","")

    with st.spinner("Step 2：抓 PubMed 中…"):
        st.session_state.pubmed_progress = st.progress(0.0)
        st.session_state.pubmed_fetch_progress = st.progress(0.0)
        dfp, cnt = fetch_pubmed(q, max_records=int(auto_pubmed_max), batch_size=200, polite_delay=0.0)
        st.session_state.pubmed_progress.empty()
        st.session_state.pubmed_fetch_progress.empty()
        st.session_state["pubmed_total_count"] = cnt

    dfs = []
    df_all_raw = pd.DataFrame()
    if not dfp.empty:
        dfs.append(dfp)
    ct_cnt = 0
    if auto_use_ctgov:
        with st.spinner("Step 2：抓 ClinicalTrials.gov 中…"):
            dft = fetch_ctgov(q, max_records=int(auto_ctgov_max))
            if not dft.empty:
                ct_cnt = len(dft)
                dfs.append(dft)

    if not dfs:
        st.error("沒有抓到任何資料。")
        return

    df_all_raw = pd.concat(dfs, ignore_index=True)
    df_all_raw = ensure_columns(df_all_raw, ["record_id","pmid","pmcid","doi","title","abstract","year",
                                             "first_author","journal","source","url","doi_url","pmc_url"], default="")
    df_dedup = deduplicate(df_all_raw)
    st.session_state["records_df"] = df_dedup
    st.session_state["screen_cursor"] = 0

    # Init per-record states
    for rid in df_dedup["record_id"].tolist():
        st.session_state["ta_final"].setdefault(rid, "Unsure")
        st.session_state["ft_decisions"].setdefault(rid, "Not reviewed yet")
        st.session_state["ft_reasons"].setdefault(rid, "")
        st.session_state["ft_text"].setdefault(rid, "")
        st.session_state["ft_note"].setdefault(rid, "")
        st.session_state["rob2"].setdefault(rid, {})

    # Choose engine
    if auto_engine == "LLM":
        use_llm = True
    elif auto_engine == "Rule-based":
        use_llm = False
    else:
        use_llm = llm_available()

    prog = st.progress(0.0)
    ai_rows = []
    with st.spinner("Step 3：AI Title/Abstract 全量粗篩中…"):
        n = len(df_dedup)
        for i in range(0, n, int(auto_batch)):
            chunk = df_dedup.iloc[i:i+int(auto_batch)].copy()
            if use_llm and llm_available():
                df_ai = run_ta_llm(chunk, pico0, criteria_now)
            else:
                df_ai = run_ta_rule(chunk, pico0)
            ai_rows.append(df_ai)
            for _, rr in df_ai.iterrows():
                st.session_state["ta_final"][rr["record_id"]] = rr["AI_label"]
            prog.progress(min(1.0, (i+len(chunk))/max(n,1)))
    prog.empty()

    st.session_state["ai_ta_df"] = pd.concat(ai_rows, ignore_index=True).drop_duplicates(subset=["record_id"], keep="last")
    summary_df0 = build_summary_df(df_dedup)
    prisma = compute_prisma_counts(df_all_raw, df_dedup)
    st.session_state["AUTO_PRISMA"] = prisma

    search_rows = auto_build_search_rows(q, st.session_state.get("pubmed_total_count",0), auto_use_ctgov, ct_cnt)

    if HAS_DOCX:
        st.session_state["AUTO_WORD_BYTES"] = make_word_supplement(q, prisma, summary_df0, search_rows)
    else:
        st.session_state["AUTO_WORD_BYTES"] = b""

    # Initialize extraction wide sheet skeleton
    base_frame = summary_df0[["record_id","pmid","doi","title","institution_openurl","FT_note_cannot_obtain"]].copy()
    for c in schema["base_cols"]:
        if c not in base_frame.columns:
            base_frame[c] = ""
    for ocol in schema["outcomes"]:
        if ocol not in base_frame.columns:
            base_frame[ocol] = ""
    for c in ["Effect_measure","Effect","Lower_CI","Upper_CI","Timepoint","Unit"]:
        if c not in base_frame.columns:
            base_frame[c] = ""
    st.session_state["extract_wide_df"] = base_frame

if st.button("一鍵跑：抓文獻→AI 粗篩→PRISMA→匯出 Word"):
    run_one_click_pipeline(
        auto_use_ctgov=auto_use_ctgov,
        auto_pubmed_max=int(auto_pubmed_max),
        auto_ctgov_max=int(auto_ctgov_max),
        auto_batch=int(auto_batch),
        auto_engine=auto_engine
    )

if st.session_state.get("AUTO_WORD_BYTES"):
    st.success("已完成一鍵流程，可下載 Word。")
    st.download_button(
        "下載 Word supplement（docx）",
        data=st.session_state["AUTO_WORD_BYTES"],
        file_name="supplement_auto.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
elif st.session_state.get("AUTO_PRISMA") and not HAS_DOCX:
    st.warning("已完成一鍵流程（含 PRISMA），但環境未安裝 python-docx，因此無法輸出 Word。")

st.markdown("---")

# =========================================================
# [12] Step 2: Fetch records (manual mode)
# =========================================================
st.header("Step 2. 抓文獻（手動模式；若你用一鍵跑可略過）")

use_pubmed = st.checkbox("PubMed", value=True)
use_ctgov = st.checkbox("ClinicalTrials.gov（可選）", value=False)

colLim1, colLim2 = st.columns([1,1])
with colLim1:
    pubmed_max = st.number_input("PubMed 抓取上限（0=全部；很大會很慢）", min_value=0, max_value=2000000, value=0, step=200)
with colLim2:
    ctgov_max = st.number_input("CT.gov 抓取上限", min_value=0, max_value=20000, value=200, step=50)

polite_delay = st.slider("（可選）API 友善延遲（秒）", min_value=0.0, max_value=1.0, value=0.0, step=0.1)

if st.button("開始抓文獻（Step 2）"):
    q = (st.session_state.get("PUBMED_QUERY","") or "").strip()
    if not q:
        st.error("PubMed 搜尋式為空：請至少填 P 或 I（或手動貼搜尋式）。")
    else:
        with st.spinner("抓取中…"):
            dfs = []
            total_count = 0
            try:
                if use_pubmed:
                    st.session_state.pubmed_progress = st.progress(0.0)
                    st.session_state.pubmed_fetch_progress = st.progress(0.0)
                    dfp, cnt = fetch_pubmed(q, max_records=int(pubmed_max), batch_size=200, polite_delay=float(polite_delay))
                    st.session_state.pubmed_progress.empty()
                    st.session_state.pubmed_fetch_progress.empty()
                    total_count = cnt
                    if not dfp.empty:
                        dfs.append(dfp)
                        st.success(f"PubMed 回傳：{len(dfp)}（總量估計 count≈{cnt}）")
                    else:
                        st.warning(f"PubMed：0（count≈{cnt}）")

                if use_ctgov and ctgov_max > 0:
                    dft = fetch_ctgov(q, max_records=int(ctgov_max))
                    if not dft.empty:
                        dfs.append(dft)
                        st.success(f"ClinicalTrials.gov：{len(dft)}")
                    else:
                        st.warning("ClinicalTrials.gov：0")

                if not dfs:
                    st.session_state["records_df"] = pd.DataFrame()
                    st.session_state["pubmed_total_count"] = total_count
                    st.error("沒有抓到任何資料。")
                else:
                    df_all = pd.concat(dfs, ignore_index=True)
                    df_all = ensure_columns(df_all, ["record_id","pmid","pmcid","doi","title","abstract","year",
                                                     "first_author","journal","source","url","doi_url","pmc_url"], default="")
                    identified = len(df_all)
                    df_dedup = deduplicate(df_all)
                    st.session_state["records_df"] = df_dedup
                    st.session_state["pubmed_total_count"] = total_count
                    st.session_state["screen_cursor"] = 0

                    for rid in df_dedup["record_id"].tolist():
                        st.session_state["ta_final"].setdefault(rid, "Unsure")
                        st.session_state["ft_decisions"].setdefault(rid, "Not reviewed yet")
                        st.session_state["ft_reasons"].setdefault(rid, "")
                        st.session_state["ft_text"].setdefault(rid, "")
                        st.session_state["ft_note"].setdefault(rid, "")
                        st.session_state["rob2"].setdefault(rid, {})

                    st.success(f"合併後 Identified={identified}；去重後={len(df_dedup)}")

            except Exception as e:
                st.session_state["records_df"] = pd.DataFrame()
                st.error(f"抓取失敗：{e}")

df_records = st.session_state.get("records_df", pd.DataFrame())
if df_records.empty:
    st.info("尚未抓取文獻。完成 Step 2 後才會顯示 Step 3-6。")
    st.stop()

st.caption(f"PubMed 原始 count 估計：{st.session_state.get('pubmed_total_count',0)}")

# =========================================================
# [13] Step 3: AI Title/Abstract screening
# =========================================================
st.header("Step 3. Title/Abstract AI Screening（可全自動；不需逐篇人工勾選）")

pico = {
    "P": st.session_state.get("P",""),
    "I": st.session_state.get("I",""),
    "C": st.session_state.get("C",""),
    "O": st.session_state.get("O",""),
    "X": st.session_state.get("X",""),
}

engine = st.selectbox(
    "篩選引擎",
    options=["Rule-based (fast, high recall)", "LLM (more accurate)"] if llm_available() else ["Rule-based (fast, high recall)"],
    index=0
)

batch_size = st.number_input("每次處理筆數（大量時建議 50~200）", min_value=10, max_value=500, value=100, step=10)
run_only_unscreened = st.checkbox("只跑尚未被 AI 判斷的資料", value=True)

already = set()
if isinstance(st.session_state.get("ai_ta_df"), pd.DataFrame) and not st.session_state["ai_ta_df"].empty:
    already = set(st.session_state["ai_ta_df"]["record_id"].tolist())

to_process = df_records.copy()
if run_only_unscreened:
    to_process = to_process[~to_process["record_id"].isin(already)]

cursor = st.session_state.get("screen_cursor", 0)
chunk = to_process.iloc[cursor:cursor+int(batch_size)].copy()

colS1, colS2 = st.columns([1,1])
with colS1:
    if st.button("AI 初篩：跑下一批"):
        if chunk.empty:
            st.info("沒有可處理資料（可能已全部處理完）。")
        else:
            with st.spinner(f"處理 {len(chunk)} 篇…"):
                if engine.startswith("LLM") and llm_available():
                    df_ai = run_ta_llm(chunk, pico, st.session_state.get("criteria_text",""))
                else:
                    df_ai = run_ta_rule(chunk, pico)

                old = st.session_state.get("ai_ta_df", pd.DataFrame())
                if old is None or old.empty:
                    st.session_state["ai_ta_df"] = df_ai
                else:
                    st.session_state["ai_ta_df"] = pd.concat([old, df_ai], ignore_index=True).drop_duplicates(subset=["record_id"], keep="last")

                for _, rr in df_ai.iterrows():
                    st.session_state["ta_final"][rr["record_id"]] = rr["AI_label"]

                st.session_state["screen_cursor"] = cursor + len(chunk)
                st.success("完成。")

with colS2:
    if st.button("AI 初篩：從頭重跑（清空 AI 結果）"):
        st.session_state["ai_ta_df"] = pd.DataFrame()
        st.session_state["ta_final"] = {rid: "Unsure" for rid in df_records["record_id"].tolist()}
        st.session_state["screen_cursor"] = 0
        st.success("已清空。")

ai_ta_df = st.session_state.get("ai_ta_df", pd.DataFrame())
if ai_ta_df is None or ai_ta_df.empty:
    st.warning("尚未執行 AI 初篩。請按『跑下一批』或用 Step 0 一鍵跑。")
    st.stop()

view_df = df_records.merge(ai_ta_df, on="record_id", how="left")
view_df = ensure_columns(view_df, ["AI_label","AI_reason","AI_confidence","AI_matched_rules"], default="")

ta_vals = [st.session_state["ta_final"].get(rid, "Unsure") for rid in df_records["record_id"].tolist()]
k_include = sum(1 for x in ta_vals if x == "Include")
k_exclude = sum(1 for x in ta_vals if x == "Exclude")
k_unsure  = sum(1 for x in ta_vals if x == "Unsure")

c1,c2,c3,c4 = st.columns(4)
with c1: st.markdown(f'<div class="kpi"><div class="label">AI screened</div><div class="value">{len(ai_ta_df)}</div></div>', unsafe_allow_html=True)
with c2: st.markdown(f'<div class="kpi"><div class="label">Include</div><div class="value">{k_include}</div></div>', unsafe_allow_html=True)
with c3: st.markdown(f'<div class="kpi"><div class="label">Exclude</div><div class="value">{k_exclude}</div></div>', unsafe_allow_html=True)
with c4: st.markdown(f'<div class="kpi"><div class="label">Unsure</div><div class="value">{k_unsure}</div></div>', unsafe_allow_html=True)

filter_mode = st.radio("檢視清單", ["只看 Unsure", "只看 Include", "只看 Exclude", "全部"], horizontal=True, index=0)

def want(dec: str) -> bool:
    if filter_mode == "全部": return True
    if filter_mode == "只看 Unsure": return dec == "Unsure"
    if filter_mode == "只看 Include": return dec == "Include"
    if filter_mode == "只看 Exclude": return dec == "Exclude"
    return True

st.markdown('<hr class="soft">', unsafe_allow_html=True)
st.subheader("AI 篩選結果（可稽核）：點開看每篇理由 + 全文欄位")

for _, row in view_df.iterrows():
    rid = row["record_id"]
    ta_dec = st.session_state["ta_final"].get(rid, "Unsure")
    if not want(ta_dec):
        continue

    title = row.get("title","") or rid
    pmid = row.get("pmid","")
    doi  = row.get("doi","")
    pmcid= row.get("pmcid","")
    year = row.get("year","")
    fa   = row.get("first_author","")
    url  = row.get("url","")
    doi_url = row.get("doi_url","")
    pmc_link= row.get("pmc_url","")

    openurl = build_openurl(st.session_state.get("RESOLVER_BASE",""), doi=doi, pmid=pmid, title=title)
    openurl = apply_ezproxy(st.session_state.get("EZPROXY",""), openurl) if openurl else ""
    pub_link = apply_ezproxy(st.session_state.get("EZPROXY",""), url) if url else ""
    doi_link = apply_ezproxy(st.session_state.get("EZPROXY",""), doi_url) if doi_url else ""
    pmc_link2= apply_ezproxy(st.session_state.get("EZPROXY",""), pmc_link) if pmc_link else ""

    with st.expander(title, expanded=False):
        st.markdown('<div class="card">', unsafe_allow_html=True)

        meta = f"<div class='meta'><b>ID</b>: {rid}"
        if pmid: meta += f" &nbsp;&nbsp; <b>PMID</b>: {pmid}"
        if doi:  meta += f" &nbsp;&nbsp; <b>DOI</b>: {doi}"
        if year: meta += f" &nbsp;&nbsp; <b>Year</b>: {year}"
        if fa:   meta += f" &nbsp;&nbsp; <b>First author</b>: {fa}"
        meta += f" &nbsp;&nbsp; <b>Source</b>: {row.get('source','')}"
        meta += "</div>"
        st.markdown(meta, unsafe_allow_html=True)

        links = []
        if pub_link: links.append(f"[PubMed/Link]({pub_link})")
        if doi_link: links.append(f"[DOI]({doi_link})")
        if pmc_link2: links.append(f"[PMC OA]({pmc_link2})")
        if openurl: links.append(f"[全文(OpenURL)]({openurl})")
        if links:
            st.markdown(" | ".join(links))

        st.markdown(badge_html(ta_dec) + f"<span class='small'>AI Title/Abstract 決策</span>", unsafe_allow_html=True)
        st.write(f"理由：{row.get('AI_reason','')}")
        if row.get("AI_matched_rules",""):
            st.caption(f"Matched rules: {row.get('AI_matched_rules','')}")
        st.caption(f"信心度：{row.get('AI_confidence','')}")

        st.markdown("### Abstract")
        st.write(row.get("abstract","") or "_No abstract available._")

        st.markdown('<hr class="soft">', unsafe_allow_html=True)

        st.markdown("### Full-text decision（看完全文後回填）")
        ft_opts = ["Not reviewed yet", "Include for meta-analysis", "Exclude after full-text"]
        cur_ft = st.session_state["ft_decisions"].get(rid, "Not reviewed yet")
        if cur_ft not in ft_opts:
            cur_ft = "Not reviewed yet"
        new_ft = st.radio("", ft_opts, index=ft_opts.index(cur_ft), key=f"ft_{rid}")
        st.session_state["ft_decisions"][rid] = new_ft

        ft_reason = st.text_area("Full-text reason / notes",
                                 value=st.session_state["ft_reasons"].get(rid,""),
                                 key=f"ft_reason_{rid}", height=80)
        st.session_state["ft_reasons"][rid] = ft_reason

        ft_note = st.text_input("若查不到全文：填原因/狀態（付費牆、館際、等待作者回信…）",
                                value=st.session_state["ft_note"].get(rid,""),
                                key=f"ft_note_{rid}")
        st.session_state["ft_note"][rid] = ft_note

        st.markdown("#### 上傳 PDF（可選）")
        uploaded_pdf = st.file_uploader("PDF 上傳（每篇文章各自上傳）", type=["pdf"], key=f"pdf_{rid}")
        extracted = ""
        if uploaded_pdf is not None:
            if HAS_PYPDF2:
                try:
                    reader = PdfReader(uploaded_pdf)
                    texts = []
                    for page in reader.pages[:60]:
                        t = page.extract_text() or ""
                        if t.strip():
                            texts.append(t)
                    extracted = "\n".join(texts).strip()
                    if not extracted:
                        st.warning("PDF 可能是掃描圖檔或無文字層。建議先 OCR（Adobe/Drive OCR）再上傳，或貼 figure/table 相關段落。")
                    else:
                        st.success(f"已抽取文字（前 60 頁），長度={len(extracted)}。")
                except Exception as e:
                    st.error(f"PDF 讀取失敗：{e}")
            else:
                st.warning("環境無 PyPDF2，無法從 PDF 抽字。請改用貼全文。")

        st.markdown("#### Full text / 關鍵段落（可貼全文、或貼 figure/table 相關段落）")
        default_text = st.session_state["ft_text"].get(rid,"")
        if extracted and len(extracted) > len(default_text):
            default_text = extracted
        ft_text = st.text_area("", value=default_text, key=f"ft_text_{rid}", height=180)
        st.session_state["ft_text"][rid] = ft_text

        st.markdown("</div>", unsafe_allow_html=True)

# =========================================================
# [15] Step 4: PRISMA + exports
# =========================================================
st.header("Step 4. PRISMA（AI 篩選）+ 匯出")

summary_df = build_summary_df(df_records)

tab1, tab2, tab3 = st.tabs(["Unsure（待確認 full-text）", "AI Excluded（含理由）", "AI Included（進 full-text）"])
with tab1:
    dfu = summary_df[summary_df["TA_final_AI"]=="Unsure"].copy()
    st.dataframe(dfu[["record_id","pmid","doi","year","first_author","title","AI_reason","institution_openurl"]], use_container_width=True)
    st.download_button("下載 Unsure 清單（CSV）", data=to_csv_bytes(dfu), file_name="ta_unsure.csv", mime="text/csv")
with tab2:
    dfx = summary_df[summary_df["TA_final_AI"]=="Exclude"].copy()
    st.dataframe(dfx[["record_id","pmid","doi","year","first_author","title","AI_reason","institution_openurl"]], use_container_width=True)
    st.download_button("下載 Excluded 清單（CSV, 含理由）", data=to_csv_bytes(dfx), file_name="ta_excluded.csv", mime="text/csv")
with tab3:
    dfi = summary_df[summary_df["TA_final_AI"]=="Include"].copy()
    st.dataframe(dfi[["record_id","pmid","doi","year","first_author","title","AI_reason","institution_openurl"]], use_container_width=True)
    st.download_button("下載 Included(for FT) 清單（CSV）", data=to_csv_bytes(dfi), file_name="ta_included.csv", mime="text/csv")

st.download_button("下載完整 screening summary（CSV）", data=to_csv_bytes(summary_df), file_name="screening_summary_all.csv", mime="text/csv")

# =========================================================
# [17] Step 5: Extraction wide sheet + optional AI extraction
# =========================================================
st.header("Step 5. Data extraction（寬表；欄位不寫死）")

base_frame = summary_df[["record_id","pmid","doi","title","institution_openurl","FT_note_cannot_obtain"]].copy()
for c in schema["base_cols"]:
    if c not in base_frame.columns:
        base_frame[c] = ""
for ocol in schema["outcomes"]:
    if ocol not in base_frame.columns:
        base_frame[ocol] = ""
for c in ["Effect_measure","Effect","Lower_CI","Upper_CI","Timepoint","Unit"]:
    if c not in base_frame.columns:
        base_frame[c] = ""

existing = st.session_state.get("extract_wide_df", pd.DataFrame())
if isinstance(existing, pd.DataFrame) and not existing.empty and "record_id" in existing.columns:
    keep_cols = [c for c in existing.columns if c in base_frame.columns] + [c for c in base_frame.columns if c not in existing.columns]
    existing = existing.reindex(columns=keep_cols)
    base_frame = base_frame.merge(existing.drop_duplicates(subset=["record_id"]), on="record_id", how="left", suffixes=("","_old"))
    for c in list(base_frame.columns):
        if c.endswith("_old"):
            orig = c[:-4]
            base_frame[orig] = base_frame[orig].astype(str)
            base_frame[c] = base_frame[c].astype(str)
            base_frame[orig] = base_frame.apply(lambda r: r[c] if (r[c].strip() != "" and r[c] != "nan") else r[orig], axis=1)
            base_frame = base_frame.drop(columns=[c])

edited = st.data_editor(
    base_frame,
    use_container_width=True,
    num_rows="dynamic",
    hide_index=True,
    column_config={
        "record_id": st.column_config.TextColumn("record_id", disabled=True),
        "pmid": st.column_config.TextColumn("PMID", disabled=True),
        "doi": st.column_config.TextColumn("DOI", disabled=True),
        "title": st.column_config.TextColumn("Title", disabled=True, width="large"),
        "institution_openurl": st.column_config.LinkColumn("全文(openurl)", display_text="open"),
        "FT_note_cannot_obtain": st.column_config.TextColumn("Fulltext note"),
        "Effect_measure": st.column_config.SelectboxColumn("Effect measure", options=["","OR","RR","HR","MD","SMD","Risk difference","Other"]),
        "Effect": st.column_config.NumberColumn("Effect", required=False, format="%.4f"),
        "Lower_CI": st.column_config.NumberColumn("Lower 95% CI", required=False, format="%.4f"),
        "Upper_CI": st.column_config.NumberColumn("Upper 95% CI", required=False, format="%.4f"),
    }
)
st.session_state["extract_wide_df"] = edited
st.download_button("下載 extraction 寬表（CSV）", data=to_csv_bytes(edited), file_name="extraction_wide.csv", mime="text/csv")

st.markdown("---")
st.subheader("5B.（可選）AI extraction（含 OCR/figure/table 提示；抽不到留空）")

if llm_available():
    n_ai = st.number_input("每次 AI 抽取筆數", min_value=1, max_value=30, value=5, step=1)
    if st.button("執行 AI extraction（對前 N 筆已有全文者）"):
        with st.spinner("AI 抽取中…"):
            text_map = st.session_state["ft_text"]
            targets = []
            for rid in edited["record_id"].tolist():
                t = (text_map.get(rid) or "").strip()
                if t:
                    targets.append(rid)
                if len(targets) >= int(n_ai):
                    break

            if not targets:
                st.warning("沒有找到已貼/已抽取全文的研究。")
            else:
                prompt_template = build_extraction_prompt(schema, pico, st.session_state.get("criteria_text",""))
                for rid in targets:
                    fulltext = (text_map.get(rid) or "").strip()
                    messages = [
                        {"role":"system","content":"You are an SR/MA full-text reviewer and extractor. Output ONLY valid JSON."},
                        {"role":"user","content":prompt_template + "\n\n[Full text]\n" + fulltext[:120000]}
                    ]
                    txt = llm_chat(messages, temperature=0.1, timeout=120)
                    js = json_from_text(txt or "")
                    if not js:
                        continue

                    d = js.get("fulltext_decision")
                    if d in ["Include for meta-analysis","Exclude after full-text","Not reviewed yet"]:
                        st.session_state["ft_decisions"][rid] = d
                    if js.get("fulltext_reason"):
                        st.session_state["ft_reasons"][rid] = str(js.get("fulltext_reason"))

                    fields = js.get("extracted_fields") or {}
                    for k,v in fields.items():
                        if k in edited.columns:
                            edited.loc[edited["record_id"]==rid, k] = str(v)

                    meta = js.get("meta") or {}
                    if "effect_measure" in meta and "Effect_measure" in edited.columns:
                        edited.loc[edited["record_id"]==rid, "Effect_measure"] = str(meta.get("effect_measure",""))
                    if "effect" in meta and "Effect" in edited.columns:
                        edited.loc[edited["record_id"]==rid, "Effect"] = meta.get("effect","")
                    if "lower_CI" in meta and "Lower_CI" in edited.columns:
                        edited.loc[edited["record_id"]==rid, "Lower_CI"] = meta.get("lower_CI","")
                    if "upper_CI" in meta and "Upper_CI" in edited.columns:
                        edited.loc[edited["record_id"]==rid, "Upper_CI"] = meta.get("upper_CI","")
                    if "timepoint" in meta and "Timepoint" in edited.columns:
                        edited.loc[edited["record_id"]==rid, "Timepoint"] = str(meta.get("timepoint",""))
                    if "unit" in meta and "Unit" in edited.columns:
                        edited.loc[edited["record_id"]==rid, "Unit"] = str(meta.get("unit",""))

                st.session_state["extract_wide_df"] = edited
                st.success("AI 抽取完成（已回填到寬表）。")
else:
    st.info("未設定 LLM：可先用寬表手動抽取。")

st.markdown("---")
st.subheader("5C. Forest plot（fixed effect；使用寬表 Effect/CI）")

if not HAS_MPL:
    st.info("環境缺少 matplotlib，跳過森林圖。")
else:
    df_fx = st.session_state.get("extract_wide_df", pd.DataFrame()).copy()
    if df_fx.empty:
        st.info("尚無 extraction 寬表。")
    else:
        measure = st.selectbox("Effect measure", options=["OR","RR","HR","MD","SMD","Risk difference"], index=0)
        sub = df_fx[df_fx["Effect_measure"].astype(str)==measure].copy()
        sub = sub.dropna(subset=["Effect","Lower_CI","Upper_CI"], how="any")
        if sub.empty:
            st.info("沒有可用的 effect/CI（請先在寬表填 Effect/CI 與 measure）。")
        else:
            try:
                eff = sub["Effect"].astype(float).tolist()
                lcl = sub["Lower_CI"].astype(float).tolist()
                ucl = sub["Upper_CI"].astype(float).tolist()

                if measure in ["OR","RR","HR"]:
                    log_eff = [math.log(x) for x in eff]
                    log_lcl = [math.log(x) for x in lcl]
                    log_ucl = [math.log(x) for x in ucl]
                    se = [(hi - lo) / (2*1.96) for lo, hi in zip(log_lcl, log_ucl)]
                    w = [1/(s*s) if s and s>0 else 0 for s in se]
                    sw = sum(w)
                    pooled_log = sum(wi*xi for wi,xi in zip(w,log_eff))/sw
                    se_pool = math.sqrt(1/sw)
                    pooled = math.exp(pooled_log)
                    pooled_l = math.exp(pooled_log - 1.96*se_pool)
                    pooled_u = math.exp(pooled_log + 1.96*se_pool)
                    null = 1.0
                else:
                    se = [(hi - lo) / (2*1.96) for lo, hi in zip(lcl, ucl)]
                    w = [1/(s*s) if s and s>0 else 0 for s in se]
                    sw = sum(w)
                    pooled = sum(wi*xi for wi,xi in zip(w,eff))/sw
                    se_pool = math.sqrt(1/sw)
                    pooled_l = pooled - 1.96*se_pool
                    pooled_u = pooled + 1.96*se_pool
                    null = 0.0

                st.write(f"Fixed-effect pooled {measure} = **{pooled:.4f}** (95% CI {pooled_l:.4f}–{pooled_u:.4f})")

                labels = []
                for _, r in sub.iterrows():
                    lab = f"{r.get('First author','') or r.get('first_author','')} {r.get('Year','') or r.get('year','')}".strip()
                    labels.append(short(lab or r.get("record_id",""), 40))

                y = list(range(len(sub), 0, -1))
                fig_h = max(3.5, 0.35*len(sub) + 2.0)
                fig, ax = plt.subplots(figsize=(8, fig_h))
                ax.errorbar(
                    eff, y,
                    xerr=[[e-l for e,l in zip(eff,lcl)], [u-e for u,e in zip(ucl,eff)]],
                    fmt='o', capsize=3
                )
                ax.axvline(null, linestyle='--', linewidth=1)
                ax.axvline(pooled, linestyle='-', linewidth=1)
                ax.set_yticks(y)
                ax.set_yticklabels(labels)
                ax.set_xlabel(f"Effect ({measure})")
                ax.set_title(f"Forest plot (fixed effect) - {measure}")
                ax.set_ylim(0, len(sub)+1)
                st.pyplot(fig)

            except Exception as e:
                st.error(f"森林圖計算失敗：{e}")

# =========================================================
# Step 6: ROB 2.0 (manual)
# =========================================================
st.header("Step 6. ROB 2.0（手動下拉）")
st.caption("ROB 2.0 需要標準化評分；這裡提供五大 domain + overall 的下拉。")

rob_candidates = summary_df[(summary_df["FT_decision"]=="Include for meta-analysis")].copy()
if rob_candidates.empty:
    st.info("目前沒有 FT=Include for meta-analysis 的研究；ROB 2.0 通常在納入後做。")
else:
    rob_levels = ["", "Low risk", "Some concerns", "High risk"]
    domain_labels = [("D1","Randomization process"),
                     ("D2","Deviations from intended interventions"),
                     ("D3","Missing outcome data"),
                     ("D4","Measurement of the outcome"),
                     ("D5","Selection of the reported result"),
                     ("Overall","Overall Risk of Bias")]
    for _, r in rob_candidates.iterrows():
        rid = r["record_id"]
        st.markdown(f"**{r.get('first_author','')} ({r.get('year','')})** — {short(r.get('title',''), 120)}")
        cols = st.columns(6)
        rb = st.session_state["rob2"].get(rid, {}) or {}
        for i,(k,lab) in enumerate(domain_labels):
            with cols[i]:
                val = st.selectbox(lab, options=rob_levels,
                                   index=rob_levels.index(rb.get(k,"")) if rb.get(k,"") in rob_levels else 0,
                                   key=f"rob_{rid}_{k}")
                rb[k] = val
        st.session_state["rob2"][rid] = rb
        st.markdown("---")

    out_rows = []
    for _, r in rob_candidates.iterrows():
        rid = r["record_id"]
        name = f"{r.get('first_author','')} ({r.get('year','')})".strip() or rid
        rb = st.session_state["rob2"].get(rid, {}) or {}
        out_rows.append({
            "Study Name": name,
            "D1 Randomization": rb.get("D1",""),
            "D2 Deviations": rb.get("D2",""),
            "D3 Missing data": rb.get("D3",""),
            "D4 Measurement": rb.get("D4",""),
            "D5 Reporting": rb.get("D5",""),
            "Overall": rb.get("Overall",""),
        })
    df_rob = pd.DataFrame(out_rows)
    st.download_button("下載 ROB2（CSV）", data=to_csv_bytes(df_rob), file_name="rob2.csv", mime="text/csv")
